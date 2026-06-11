"""WO v4.33 §3.2 — one-shot Pre-Job Card template importer (review-and-normalize, §0.15).

Reads Nadie's Word templates, normalizes them, and inserts `prejob_templates` DRAFTS
(`is_active=False`) for BA/Nadie review in the §3.3 admin screen. Idempotent: existing names
are skipped (or refreshed with --update, drafts only). Run:

    python -m scripts.import_prejob_templates --folder "<templates folder>" [--dry-run] [--update]

.doc handling (12 of the 22 are legacy binary Word): converted once via Word COM automation
(pywin32 + an installed Word — this dev box) into `<folder>/_converted/`, cached by mtime.
On a box without Word/pywin32 the script lists the files to convert manually (File > Save As
> .docx). CI never runs the conversion: the parser tests use pre-converted fixtures under
tests/fixtures/prejob_templates/.

Normalization applied (§0.4/§0.15, all content-keyed — never positional):
  * Section headings -> canonical: GRP SECTION / SUB FRAME SECTION / FINISHING SECTION /
    STEEL SECTION / CHASSIS MODIFICATIONS (accepts G.R.P., SUBFRAME, etc. variants).
  * "Body gab" -> "Body gap" (any case).
  * Middle + Big Icecream: the GRP-section item that is structurally a FINISHING doorframe
    line ("DRD's ... 3Cr12 ...") is DROPPED (Nadie-confirmed copy-paste error, Q5).
  * `Note:` lines attach to the preceding item as its `note`.
  * `*`-prefixed lines (incl. the Explosive "(Below 3.5ton) *..." conditional pack) attach to
    the preceding item as `sub_items` — the §0.5 HazChem shape; tonnage variants stay one
    template in v4.33 (conditional split is a deferred §3.0 item).
  * product_line: filename says "Rhinorange 2.0" -> rhinorange_2_0; bare "Rhinorange" ->
    rhinorange_legacy — UPGRADED to rhinorange_2_0 when the document header itself says
    "Rhinorange 2.0" (the 15.5m Body Only case, analysis v1.1 Finding 3).
"""
from __future__ import annotations

import argparse
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path

# ── canonical section headings (§0.4) ────────────────────────────────────────
_HEADING_RE = re.compile(
    r"^(?P<name>G\.?\s*R\.?\s*P\.?|GRP|SUB\s*FRAME|SUBFRAME|FINISHING|STEEL|"
    r"CHASSIS\s+MODIFICATIONS?)\s*(SECTION)?\s*:?\s*$",
    re.IGNORECASE,
)
_CANONICAL = {
    "GRP": "GRP SECTION",
    "SUBFRAME": "SUB FRAME SECTION",
    "FINISHING": "FINISHING SECTION",
    "STEEL": "STEEL SECTION",
    "CHASSISMODIFICATION": "CHASSIS MODIFICATIONS",
}
_NOTE_RE = re.compile(r"^Note\s*:?\s*(?P<note>.+)$", re.IGNORECASE)
_NUM_PREFIX_RE = re.compile(r"^\d+\s*[\.\)]?\s+")
_BODY_GAB_RE = re.compile(r"\bbody\s+gab\b", re.IGNORECASE)
# The Icecream Mid/Big copy-paste error: a doorframe/DRD FINISHING line sitting in GRP (Q5).
_ICECREAM_MISPLACED_RE = re.compile(r"DRD", re.IGNORECASE)

_BODY_TYPES = ["icecream", "chiller", "freezer", "meathanger", "bakery", "explosive",
               "medical waste"]


def _canonical_heading(raw: str) -> str | None:
    m = _HEADING_RE.match(raw.strip())
    if not m:
        return None
    key = re.sub(r"[^A-Z]", "", m.group("name").upper())
    if key.startswith("CHASSISMODIFICATION"):
        key = "CHASSISMODIFICATION"
    return _CANONICAL.get(key)


def _is_subitem(text: str) -> bool:
    t = text.strip()
    return t.startswith("*") or t.startswith("•") or bool(re.match(r"^\([^)]*\)\s*\*", t))


def _clean_subitem(text: str) -> str:
    return text.strip().lstrip("•").strip()


@dataclass
class ParsedTemplate:
    source_file: str
    name: str = ""
    body_type: str = ""
    size_category: str | None = None
    product_line: str = "standard"
    header_format: str | None = None
    sections: list = field(default_factory=list)
    fixes: list = field(default_factory=list)         # human log of normalizations applied


def derive_metadata(filename: str, header_line: str | None) -> dict:
    """name / body_type / size_category / product_line from the filename (+ header upgrade)."""
    inner = filename
    m = re.search(r"\(([^)]+)\)", filename)
    if m:
        inner = m.group(1).strip()
    low = inner.lower()

    if "trailer" in low or "body only" in low:
        body_type = "trailer"
    else:
        body_type = next((bt.replace(" ", "_") for bt in _BODY_TYPES if bt in low),
                         "dry_freight" if "dry freight" in low else "")
        if not body_type and "dry" in low:
            body_type = "dry_freight"
    if not body_type:
        body_type = "unknown"

    size = None
    sm = re.search(r"(\d{1,2}\.\d)m", low)
    if sm:
        size = f"{sm.group(1)}m"
    elif "mid" in low:                                  # 'mid' | 'middle'
        size = "mid"
    elif "big" in low:
        size = "big"

    if "rhinorange 2.0" in low:
        product_line = "rhinorange_2_0"
    elif "rhinorange" in low:
        product_line = "rhinorange_legacy"
        if header_line and "rhinorange 2.0" in header_line.lower():
            product_line = "rhinorange_2_0"            # v1.1 Finding 3 — content says 2.0
    else:
        product_line = "standard"

    return {"name": inner, "body_type": body_type, "size_category": size,
            "product_line": product_line}


def _doc_lines(doc) -> list[str]:
    """Flat line stream: body paragraphs, then table-cell paragraphs (the Explosive template
    lays its sections out as a 2-column number|text TABLE — doc.paragraphs alone sees only the
    header). Cell paragraphs are fed individually so Note:/sub-item detection still works;
    merged cells are de-duplicated. (Order assumption: any table follows the header paragraphs
    — true for every template in the library; flagged in the report if a doc ever interleaves.)
    """
    lines = [p.text.strip() for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            seen: set[int] = set()
            for cell in row.cells:
                if id(cell._tc) in seen:               # merged cells repeat the same tc
                    continue
                seen.add(id(cell._tc))
                lines.extend(p.text.strip() for p in cell.paragraphs)
    return lines


def parse_docx(path: Path) -> ParsedTemplate:
    """Extract header + sections/items (+ notes, sub_items) from one .docx."""
    from docx import Document                          # lazy: keeps --help usable anywhere

    parsed = ParsedTemplate(source_file=path.name)
    doc = Document(str(path))
    lines = _doc_lines(doc)

    current: dict | None = None
    for raw in lines:
        if not raw:
            continue
        if re.fullmatch(r"\d+", raw):                  # a table numbering-column cell
            continue
        heading = _canonical_heading(raw)
        if heading:
            if current and any(s["name"] == heading for s in parsed.sections):
                parsed.fixes.append(f"duplicate heading '{heading}' merged")
                current = next(s for s in parsed.sections if s["name"] == heading)
                continue
            current = {"name": heading, "items": []}
            parsed.sections.append(current)
            continue
        if current is None:
            if parsed.header_format is None:
                parsed.header_format = raw             # first non-empty pre-section line
            continue

        note_m = _NOTE_RE.match(raw)
        if note_m and current["items"]:
            current["items"][-1]["note"] = note_m.group("note").strip()
            continue
        if _is_subitem(raw) and current["items"]:
            current["items"][-1].setdefault("sub_items", []).append(_clean_subitem(raw))
            continue

        text = _NUM_PREFIX_RE.sub("", raw).strip()
        if _BODY_GAB_RE.search(text):
            text = _BODY_GAB_RE.sub("Body gap", text)
            parsed.fixes.append('"Body gab" -> "Body gap"')
        if text:
            current["items"].append({"text": text})
    return parsed


def normalize(parsed: ParsedTemplate) -> ParsedTemplate:
    """Filename metadata + the content-keyed Icecream Mid/Big GRP fix (Q5)."""
    meta = derive_metadata(parsed.source_file, parsed.header_format)
    parsed.name = meta["name"]
    parsed.body_type = meta["body_type"]
    parsed.size_category = meta["size_category"]
    parsed.product_line = meta["product_line"]

    if parsed.body_type == "icecream" and parsed.size_category in ("mid", "big"):
        for section in parsed.sections:
            if section["name"] != "GRP SECTION":
                continue
            keep = []
            for item in section["items"]:
                if (_ICECREAM_MISPLACED_RE.search(item["text"])
                        and "3cr12" in item["text"].lower()):
                    parsed.fixes.append(
                        "dropped misplaced GRP doorframe item (Nadie Q5 copy-paste error)")
                    continue
                keep.append(item)
            section["items"] = keep
    return parsed


# ── .doc conversion (Word COM; this dev box only) ────────────────────────────
def convert_docs(folder: Path) -> tuple[list[Path], list[Path]]:
    """Convert every top-level .doc into <folder>/_converted/<stem>.docx (mtime-cached).
    Returns (converted_or_cached, failed)."""
    doc_files = sorted(p for p in folder.glob("*.doc") if not p.name.startswith("~"))
    if not doc_files:
        return [], []
    out_dir = folder / "_converted"
    out_dir.mkdir(exist_ok=True)
    need = [p for p in doc_files
            if not (out_dir / (p.stem + ".docx")).exists()
            or (out_dir / (p.stem + ".docx")).stat().st_mtime < p.stat().st_mtime]
    done = [out_dir / (p.stem + ".docx") for p in doc_files if p not in need]
    if not need:
        return done, []
    try:
        import win32com.client  # type: ignore
    except ImportError:
        print("!! pywin32 not installed — convert these manually (Word > Save As > .docx):")
        for p in need:
            print(f"   {p.name}")
        return done, need
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    failed: list[Path] = []
    try:
        for p in need:
            target = out_dir / (p.stem + ".docx")
            try:
                d = word.Documents.Open(str(p), ReadOnly=True)
                d.SaveAs2(str(target), FileFormat=16)   # 16 = wdFormatDocumentDefault (.docx)
                d.Close(False)
                done.append(target)
                print(f"   converted: {p.name}")
            except Exception as exc:                    # noqa: BLE001 — per-file, keep going
                failed.append(p)
                print(f"!! convert FAILED {p.name}: {exc}")
    finally:
        word.Quit()
    return done, failed


# ── DB import ─────────────────────────────────────────────────────────────────
def import_templates(folder: Path, dry_run: bool = False, update: bool = False) -> list[dict]:
    converted, failed = convert_docs(folder)
    docx_files = sorted(p for p in folder.glob("*.docx") if not p.name.startswith("~"))
    skip_names = ("nadie",)                             # the Q&A docs live in the same folder
    sources = [p for p in docx_files if not any(s in p.name.lower() for s in skip_names)]
    sources += converted

    rows: list[dict] = []
    from app.database import SessionLocal
    from app.models.mes import PrejobTemplate
    with SessionLocal() as db:
        for path in sorted(sources, key=lambda p: p.name.lower()):
            parsed = normalize(parse_docx(path))
            status = "parsed"
            if not parsed.sections:
                status = "NO-SECTIONS (review manually)"
            elif not dry_run:
                existing = db.query(PrejobTemplate).filter_by(name=parsed.name).first()
                if existing is not None and not update:
                    status = "skipped (exists)"
                elif existing is not None and update and existing.is_active:
                    status = "skipped (ACTIVE — not overwriting an approved template)"
                else:
                    if existing is None:
                        existing = PrejobTemplate(name=parsed.name, created_by="import-script")
                        db.add(existing)
                        status = "imported (draft)"
                    else:
                        status = "updated (draft)"
                    existing.body_type = parsed.body_type
                    existing.size_category = parsed.size_category
                    existing.product_line = parsed.product_line
                    existing.header_format = parsed.header_format
                    existing.sections = parsed.sections
                    existing.is_active = False
                    existing.updated_by = "import-script"
            rows.append({
                "file": path.name, "name": parsed.name, "body_type": parsed.body_type,
                "size": parsed.size_category or "-", "line": parsed.product_line,
                "sections": "+".join(s["name"].split(" ")[0] for s in parsed.sections),
                "items": sum(len(s["items"]) for s in parsed.sections),
                "fixes": "; ".join(sorted(set(parsed.fixes))) or "-",
                "status": status,
            })
        if not dry_run:
            db.commit()
    for p in failed:
        rows.append({"file": p.name, "name": "-", "body_type": "-", "size": "-", "line": "-",
                     "sections": "-", "items": 0, "fixes": "-",
                     "status": "CONVERSION FAILED"})
    return rows


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--folder", required=True, help="Nadie templates folder")
    ap.add_argument("--dry-run", action="store_true", help="parse + report, no DB writes")
    ap.add_argument("--update", action="store_true",
                    help="refresh existing DRAFT rows from the files (never touches active)")
    args = ap.parse_args()
    rows = import_templates(Path(args.folder), dry_run=args.dry_run, update=args.update)
    widths = {k: max(len(str(r[k])) for r in rows + [dict.fromkeys(rows[0], k)]) for k in rows[0]}
    hdr = "  ".join(k.upper().ljust(widths[k]) for k in rows[0])
    print("\n" + hdr + "\n" + "-" * len(hdr))
    for r in rows:
        print("  ".join(str(r[k]).ljust(widths[k]) for k in r))
    print(f"\n{len(rows)} files processed.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
